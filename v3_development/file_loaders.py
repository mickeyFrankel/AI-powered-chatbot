"""
File Loaders Module for VectorDB Q&A System

This module provides file loading functionality for various file formats.
Following the Single Responsibility Principle and Open/Closed Principle,
each loader is responsible for one file type and new loaders can be added
without modifying existing code.

Author: Mickey Frankel
Date: 2025-10-30
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Any
import pandas as pd
import docx
import PyPDF2
from openpyxl import load_workbook

from models import Document
from config import FileExtensions, DataProcessingConfig


class BaseFileLoader(ABC):
    """
    Abstract base class for file loaders.
    
    This class defines the interface that all file loaders must implement,
    following the Open/Closed Principle - open for extension, closed for modification.
    """
    
    @abstractmethod
    def load(self, file_path: Path) -> List[Document]:
        """
        Load documents from a file.
        
        Args:
            file_path: Path to the file to load
            
        Returns:
            List of Document objects
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is invalid
        """
        pass
    
    @abstractmethod
    def supports(self, file_path: Path) -> bool:
        """
        Check if this loader supports the given file.
        
        Args:
            file_path: Path to check
            
        Returns:
            True if this loader can handle the file
        """
        pass
    
    def _fix_phone_number(self, value: Any) -> Any:
        """
        Convert scientific notation phone numbers to proper format.
        
        Args:
            value: Phone number value (may be in scientific notation)
            
        Returns:
            Formatted phone number or original value
        """
        if pd.isna(value) or value == '':
            return value
        
        value_str = str(value).strip()
        
        # Detect scientific notation (9.73E+11, 5.42e+08)
        if 'e+' in value_str.lower() or 'e-' in value_str.lower():
            try:
                num = int(float(value_str))
                result = str(num)
                
                # Israeli phone: 9-10 digits, add leading 0 if missing
                if len(result) in DataProcessingConfig.ISRAELI_PHONE_DIGITS:
                    if not result.startswith('0'):
                        result = '0' + result
                
                return result
            except (ValueError, OverflowError):
                pass
        
        return value


class CSVFileLoader(BaseFileLoader):
    """
    Loader for CSV files.
    
    This loader handles CSV files with automatic phone number formatting
    and metadata extraction.
    """
    
    def supports(self, file_path: Path) -> bool:
        """Check if file is CSV."""
        return file_path.suffix.lower() == FileExtensions.CSV
    
    def load(self, file_path: Path) -> List[Document]:
        """
        Load documents from CSV file.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            List of Document objects
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        df = pd.read_csv(file_path, encoding='utf-8')
        df = self._preprocess_dataframe(df)
        df = self._fix_phone_columns(df)
        
        return self._dataframe_to_documents(df, file_path)
    
    def _preprocess_dataframe(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Clean and optimize DataFrame.
        
        Args:
            df: Input DataFrame
            
        Returns:
            Cleaned DataFrame
        """
        # Remove completely empty columns
        df = df.dropna(axis=1, how='all')
        
        # Remove sparse columns (>95% empty), except important fields
        sparse_threshold = DataProcessingConfig.SPARSE_COLUMN_THRESHOLD
        important_patterns = DataProcessingConfig.IMPORTANT_FIELD_PATTERNS
        
        cols_to_drop = []
        for col in df.columns:
            null_ratio = df[col].isnull().sum() / len(df)
            if null_ratio > sparse_threshold:
                col_lower = col.lower()
                if not any(pattern in col_lower for pattern in important_patterns):
                    cols_to_drop.append(col)
        
        df = df.drop(columns=cols_to_drop)
        
        # Trim whitespace
        for col in df.select_dtypes(include=['object']).columns:
            df[col] = df[col].apply(
                lambda x: x.strip() if isinstance(x, str) else x
            )
        
        # Remove duplicate rows
        df = df.drop_duplicates()
        
        # Remove mostly empty rows
        row_threshold = DataProcessingConfig.EMPTY_ROW_THRESHOLD
        df = df[df.isnull().sum(axis=1) / len(df.columns) < row_threshold]
        
        return df
    
    def _fix_phone_columns(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Fix phone number formatting in all phone columns.
        
        Args:
            df: DataFrame with phone columns
            
        Returns:
            DataFrame with fixed phone numbers
        """
        phone_patterns = ['phone', 'mobile', 'tel']
        
        for col in df.columns:
            if any(pattern in col.lower() for pattern in phone_patterns):
                df[col] = df[col].apply(self._fix_phone_number)
        
        return df
    
    def _dataframe_to_documents(
        self,
        df: pd.DataFrame,
        file_path: Path
    ) -> List[Document]:
        """
        Convert DataFrame rows to Document objects.
        
        Args:
            df: DataFrame to convert
            file_path: Source file path
            
        Returns:
            List of Document objects
        """
        documents = []
        
        for idx, row in df.iterrows():
            text_parts = []
            metadata = {}
            
            for col in df.columns:
                if pd.notna(row[col]):
                    text_parts.append(f"{col}: {row[col]}")
                    metadata[col] = str(row[col])
            
            documents.append(Document(
                text=' | '.join(text_parts),
                metadata=metadata,
                source=str(file_path),
                row_id=idx
            ))
        
        return documents


class ExcelFileLoader(BaseFileLoader):
    """
    Loader for Excel files (.xlsx, .xls).
    
    This loader handles Excel files with support for multiple sheets.
    """
    
    def supports(self, file_path: Path) -> bool:
        """Check if file is Excel."""
        return file_path.suffix.lower() in FileExtensions.EXCEL
    
    def load(self, file_path: Path) -> List[Document]:
        """
        Load documents from Excel file.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            List of Document objects
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        workbook = load_workbook(file_path, read_only=True)
        documents = []
        
        for sheet_name in workbook.sheetnames:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            sheet_docs = self._process_sheet(df, file_path, sheet_name)
            documents.extend(sheet_docs)
        
        return documents
    
    def _process_sheet(
        self,
        df: pd.DataFrame,
        file_path: Path,
        sheet_name: str
    ) -> List[Document]:
        """Process a single Excel sheet."""
        # Reuse CSV loader logic
        csv_loader = CSVFileLoader()
        df = csv_loader._preprocess_dataframe(df)
        df = csv_loader._fix_phone_columns(df)
        
        documents = []
        for idx, row in df.iterrows():
            text_parts = []
            metadata = {'sheet': sheet_name}
            
            for col in df.columns:
                if pd.notna(row[col]):
                    text_parts.append(f"{col}: {row[col]}")
                    metadata[col] = str(row[col])
            
            documents.append(Document(
                text=' | '.join(text_parts),
                metadata=metadata,
                source=str(file_path),
                row_id=f"{sheet_name}_{idx}"
            ))
        
        return documents


class DocxFileLoader(BaseFileLoader):
    """
    Loader for Word documents (.docx).
    
    This loader extracts paragraphs from Word documents.
    """
    
    def supports(self, file_path: Path) -> bool:
        """Check if file is DOCX."""
        return file_path.suffix.lower() == FileExtensions.DOCX
    
    def load(self, file_path: Path) -> List[Document]:
        """
        Load documents from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of Document objects
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        doc = docx.Document(file_path)
        documents = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                documents.append(Document(
                    text=paragraph.text.strip(),
                    metadata={'paragraph_id': i},
                    source=str(file_path),
                    row_id=f"para_{i}"
                ))
        
        return documents


class PDFFileLoader(BaseFileLoader):
    """
    Loader for PDF files (.pdf).
    
    This loader extracts text from PDF pages and paragraphs.
    """
    
    def supports(self, file_path: Path) -> bool:
        """Check if file is PDF."""
        return file_path.suffix.lower() == FileExtensions.PDF
    
    def load(self, file_path: Path) -> List[Document]:
        """
        Load documents from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of Document objects
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    paragraphs = [
                        p.strip() 
                        for p in text.split('\n\n') 
                        if p.strip()
                    ]
                    
                    for para_idx, paragraph in enumerate(paragraphs):
                        documents.append(Document(
                            text=paragraph,
                            metadata={
                                'page': page_num + 1,
                                'paragraph_id': para_idx
                            },
                            source=str(file_path),
                            row_id=f"page_{page_num}_para_{para_idx}"
                        ))
        
        return documents


class TxtFileLoader(BaseFileLoader):
    """
    Loader for plain text files (.txt).
    
    This loader extracts paragraphs from plain text files.
    """
    
    def supports(self, file_path: Path) -> bool:
        """Check if file is TXT."""
        return file_path.suffix.lower() == FileExtensions.TXT
    
    def load(self, file_path: Path) -> List[Document]:
        """
        Load documents from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            List of Document objects
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        documents = []
        
        for i, paragraph in enumerate(paragraphs):
            documents.append(Document(
                text=paragraph,
                metadata={'paragraph_id': i},
                source=str(file_path),
                row_id=f"para_{i}"
            ))
        
        return documents


class FileLoaderFactory:
    """
    Factory for creating appropriate file loaders.
    
    This class implements the Factory pattern and follows the Open/Closed
    Principle - new loaders can be registered without modifying existing code.
    """
    
    def __init__(self):
        """Initialize factory with default loaders."""
        self._loaders: List[BaseFileLoader] = [
            CSVFileLoader(),
            ExcelFileLoader(),
            DocxFileLoader(),
            PDFFileLoader(),
            TxtFileLoader(),
        ]
    
    def register_loader(self, loader: BaseFileLoader) -> None:
        """
        Register a new file loader.
        
        Args:
            loader: Loader instance to register
        """
        self._loaders.append(loader)
    
    def get_loader(self, file_path: Path) -> BaseFileLoader:
        """
        Get appropriate loader for a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Appropriate file loader
            
        Raises:
            ValueError: If no loader supports the file type
        """
        for loader in self._loaders:
            if loader.supports(file_path):
                return loader
        
        raise ValueError(
            f"Unsupported file format: {file_path.suffix}. "
            f"Supported formats: {FileExtensions.all_supported()}"
        )
    
    def load_file(self, file_path: Path) -> List[Document]:
        """
        Load documents from a file using appropriate loader.
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of Document objects
        """
        loader = self.get_loader(file_path)
        return loader.load(file_path)
